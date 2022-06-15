import os
import win32com.client as wc
import docx
c_p = os.getcwd()
print(c_p)
file_name = r'我的简历.doc'
f_p = os.path.abspath(os.path.join(c_p, 'files\\'+file_name))
print(f_p)
'''
# doc文件转换为docx
word = wc.Dispatch('Word.Application')
doc = word.Documents.Open(f_p)
print(doc)
new_file_name = r'我的简历.docx'
docx_file_path = os.path.abspath(os.path.join(c_p, 'files\\'+new_file_name))
doc.SaveAs(docx_file_path, 12)
doc.Close()
word.Quit()
'''
docx_file_name = r'我的简历.docx'
docx_file_name = r'软件开发主管.docx'
docx_file_path = os.path.abspath(os.path.join(c_p, 'files\\'+docx_file_name))
docx_content = docx.Document(docx_file_path)
paragraphs = docx_content.paragraphs
print(type(paragraphs), len(paragraphs))
if paragraphs:
    for p in paragraphs:
        print('Text : ',type(p), type(p.text), p.text if p.text else len(p.text))
else:
    print('The document has no content!')
tables = docx_content.tables
print('Table : ', type(tables), len(tables))
if tables:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                print('Cell Text : ', len(cell.text), cell.text)
else:
    print('The document has no tables!')